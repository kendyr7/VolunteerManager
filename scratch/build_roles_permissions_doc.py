from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
OUTPUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUTPUT_DIR / "estructura-de-roles-y-permisos.docx"

BLACK = "000000"
TEXT = "202124"
MUTED = "5F6368"
LIGHT = "F1F3F4"
BLUE = "1A73E8"
PALE_BLUE = "E8F0FE"
PALE_YELLOW = "FEF7E0"
BORDER = "DADCE0"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_font(run, size=11, bold=False, color=TEXT, name="Arial"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def add_text(p, text, *, bold=False, color=TEXT, size=11):
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def style_paragraph(p, *, before=0, after=8, line=1.15, keep=False):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(
        p,
        before={1: 20, 2: 18, 3: 16}[level],
        after={1: 6, 2: 6, 3: 4}[level],
        line=1.0,
        keep=True,
    )
    add_text(
        p,
        text,
        size={1: 20, 2: 16, 3: 14}[level],
        color=BLACK if level < 3 else "434343",
    )
    p.style = doc.styles[f"Heading {level}"]
    return p


def add_body(doc, text, *, bold_prefix=None, after=8):
    p = doc.add_paragraph()
    style_paragraph(p, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    style_paragraph(p, after=4)
    p.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    add_text(p, text)
    return p


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = numbering.num_having_numId(style_num_id)
    abstract_id = base_num.abstractNumId.val
    new_num = numbering.add_num(abstract_id)
    new_num.add_lvlOverride(ilvl=0).add_startOverride(1)
    return new_num.numId


def add_number(doc, text, num_id):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, after=5)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id
    add_text(p, text)
    return p


def add_callout(doc, title, body, color=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_border(cell, color=color, size="0")
    p = cell.paragraphs[0]
    style_paragraph(p, after=3)
    add_text(p, title, bold=True, color=BLACK)
    p2 = cell.add_paragraph()
    style_paragraph(p2, after=0)
    add_text(p2, body, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in (
        ("Heading 1", 20, BLACK),
        ("Heading 2", 16, BLACK),
        ("Heading 3", 14, "434343"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(TEXT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")


def add_role_card(doc, title, scope, summary, items):
    add_heading(doc, title, 2)
    p = doc.add_paragraph()
    style_paragraph(p, after=5)
    add_text(p, "Alcance: ", bold=True, color=BLUE)
    add_text(p, scope)
    add_body(doc, summary, after=5)
    for item in items:
        add_bullet(doc, item)


doc = Document()
configure_styles(doc)

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.5)
section.footer_distance = Inches(0.5)

# Cover
p = doc.add_paragraph()
style_paragraph(p, before=0, after=3, line=1.0)
add_text(p, "Estructura de roles y permisos", size=26, color=BLACK)

p = doc.add_paragraph()
style_paragraph(p, after=14)
add_text(p, "Volunteer Manager · Documento para revisión", size=14, color=MUTED)

p = doc.add_paragraph()
style_paragraph(p, after=24)
add_text(p, "Borrador de definición funcional · 11 de agosto de 2026", size=10, color=MUTED)

add_callout(
    doc,
    "Decisión principal",
    "El escaneo de QR y la operación de check-in/check-out pertenecen al Coordinador de tecnología. El Coordinador de comité nunca recibe esos permisos. El Administrador conserva acceso total por definición de su rol.",
    PALE_YELLOW,
)

add_heading(doc, "Propósito", 1)
add_body(
    doc,
    "Este documento define la estructura propuesta de roles, permisos, alcance de datos y reglas de seguridad. Su objetivo es servir como referencia antes de implementar cambios en la aplicación y permitir que el equipo confirme cualquier ajuste pendiente.",
)

add_heading(doc, "Principios rectores", 1)
for item in (
    "Solo los Administradores pueden activar o desactivar permisos.",
    "El rol determina la responsabilidad principal; el alcance determina sobre qué personas y comités puede actuar.",
    "La interfaz puede ocultar una acción, pero el servidor siempre debe volver a validar la autorización.",
    "Un cambio de permisos debe reflejarse inmediatamente en la sesión de la persona afectada, sin requerir cerrar sesión.",
    "Toda acción sensible debe quedar registrada con actor, persona afectada, fecha y resultado.",
):
    add_bullet(doc, item)

add_heading(doc, "1. Roles propuestos", 1)

add_role_card(
    doc,
    "Administrador",
    "Global",
    "Tiene acceso completo a la aplicación. No existe un límite fijo de Administradores, aunque se recomienda asignar este rol únicamente a las personas que realmente necesitan control total.",
    (
        "Administra usuarios, roles, permisos, ajustes e integraciones.",
        "Puede ver y modificar información de cualquier voluntario y comité.",
        "Puede realizar todas las operaciones de asistencia, incluido QR, por ser el rol de acceso total.",
        "Es el único rol que puede cambiar permisos de otros roles o usuarios.",
    ),
)

add_role_card(
    doc,
    "Coordinador de tecnología",
    "Global para la operación del evento",
    "Atiende la recepción y se encarga completamente de los check-in y check-out. Necesita buscar, ver y editar información de voluntarios de cualquier comité.",
    (
        "Puede ver a todos los voluntarios y su información personal completa.",
        "Puede editar información personal de cualquier voluntario.",
        "Es el único tipo de Coordinador que puede escanear QR y registrar check-in/check-out.",
        "El registro de asistencia faltante y la corrección manual de horarios son permisos configurables.",
        "Importar, crear o archivar voluntarios no forma parte de su responsabilidad inicial y se propone como permiso configurable, desactivado por defecto.",
    ),
)

add_role_card(
    doc,
    "Coordinador de comité",
    "Solo su comité asignado",
    "Es el presidente o responsable de un comité. Gestiona las necesidades de sus propios voluntarios, sin intervenir en la recepción ni en la asistencia global.",
    (
        "Puede ver a los voluntarios de su comité, sus avisos y sus solicitudes.",
        "Puede reagendar turnos únicamente dentro de su comité.",
        "La edición de información personal de sus voluntarios es configurable.",
        "Nunca puede escanear QR, registrar check-in/check-out, corregir asistencia global, importar ni crear voluntarios.",
    ),
)

add_role_card(
    doc,
    "Voluntario",
    "Solo su propia información",
    "Accede a las funciones personales que la aplicación ofrezca, sin permisos administrativos ni de gestión sobre otras personas.",
    (
        "Consulta su información, turnos, avisos y solicitudes personales.",
        "No administra voluntarios, comités, asistencia, usuarios ni permisos.",
    ),
)

add_heading(doc, "2. Matriz de permisos", 1)
add_body(
    doc,
    "La siguiente matriz diferencia permisos fijos, configurables y no disponibles. “Propio comité” y “global” describen el alcance máximo permitido.",
)

matrix = [
    ("Ver lista de voluntarios", "Sí · global", "Sí · global", "Sí · propio comité", "No"),
    ("Ver información personal completa", "Sí · global", "Sí · global", "Sí · propio comité", "Solo propia"),
    ("Editar información personal", "Sí · global", "Sí · global", "Configurable · propio comité", "No"),
    ("Crear voluntarios", "Sí", "Configurable · inicio No", "No", "No"),
    ("Importar voluntarios", "Sí", "Configurable · inicio No", "No", "No"),
    ("Archivar voluntarios", "Sí", "Configurable · inicio No", "No", "No"),
    ("Asignar o reagendar turnos", "Sí · global", "Sí · global", "Sí · propio comité", "No"),
    ("Ver y gestionar solicitudes", "Sí · global", "Configurable · global", "Sí · propio comité", "Solo propias"),
    ("Publicar avisos", "Sí · global", "Configurable · global", "Sí · propio comité", "No"),
    ("Consultar reportes", "Sí · global", "Configurable · global", "Configurable · propio comité", "No"),
    ("Escanear QR", "Sí", "Sí", "No · nunca", "No"),
    ("Registrar check-in/check-out", "Sí", "Sí", "No · nunca", "No"),
    ("Registrar asistencia faltante", "Sí", "Configurable", "No", "No"),
    ("Corregir horarios manualmente", "Sí", "Configurable", "No", "No"),
    ("Administrar usuarios y roles", "Sí", "No", "No", "No"),
    ("Configurar permisos", "Sí", "No", "No", "No"),
    ("Modificar ajustes del sistema", "Sí", "No", "No", "No"),
]

table = doc.add_table(rows=1, cols=5)
widths = [2700, 1050, 1850, 2460, 1300]
headers = ("Permiso", "Admin", "Coord. tecnología", "Coord. comité", "Voluntario")
for idx, text in enumerate(headers):
    cell = table.rows[0].cells[idx]
    set_cell_shading(cell, LIGHT)
    set_cell_border(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.0)
    add_text(p, text, bold=True, size=9, color=BLACK)
set_repeat_table_header(table.rows[0])

for row_idx, row in enumerate(matrix):
    cells = table.add_row().cells
    for idx, text in enumerate(row):
        cell = cells[idx]
        if row_idx % 2:
            set_cell_shading(cell, "FAFAFA")
        if "No · nunca" in text:
            set_cell_shading(cell, PALE_YELLOW)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.05)
        add_text(p, text, size=8.5, bold=(idx == 0), color=TEXT)

set_table_geometry(table, widths)

add_callout(
    doc,
    "Regla invariable de QR",
    "Ningún interruptor puede otorgar QR o check-in/check-out al Coordinador de comité. Entre coordinadores, estas acciones son exclusivas del Coordinador de tecnología.",
    PALE_YELLOW,
)

add_heading(doc, "3. Alcance y límites", 1)

add_heading(doc, "Administrador", 2)
add_body(doc, "Alcance global sobre todos los comités, voluntarios, turnos, asistencias, reportes, usuarios y ajustes.")

add_heading(doc, "Coordinador de tecnología", 2)
add_body(doc, "Alcance global sobre voluntarios y asistencia porque trabaja en recepción. Puede localizar y atender a cualquier persona, aunque pertenezca a otro comité.")

add_heading(doc, "Coordinador de comité", 2)
add_body(doc, "Alcance restringido al comité asignado. Las consultas y cambios deben filtrar por comité tanto en la interfaz como en el servidor.")

add_heading(doc, "Voluntario", 2)
add_body(doc, "Alcance personal. Solo puede consultar o ejecutar acciones asociadas con su propia cuenta y sus propios registros.")

add_heading(doc, "4. Configuración inicial recomendada", 1)

add_heading(doc, "Coordinador de tecnología", 2)
for item in (
    "Activo: ver y editar todos los voluntarios; buscar globalmente; asignar o reagendar turnos; QR; check-in y check-out.",
    "Configurable: registrar asistencia faltante; corregir horarios manualmente; avisos; solicitudes; reportes.",
    "Configurable, inicialmente desactivado: crear, importar y archivar voluntarios.",
):
    add_bullet(doc, item)

add_heading(doc, "Coordinador de comité", 2)
for item in (
    "Activo: ver voluntarios, avisos y solicitudes de su comité; reagendar a sus propios voluntarios.",
    "Configurable: editar información personal y consultar reportes de su comité.",
    "No disponible: QR, check-in, check-out, asistencia global, corrección manual de asistencia, importación, creación y archivo de voluntarios.",
):
    add_bullet(doc, item)

add_callout(
    doc,
    "Control administrativo",
    "Los interruptores de permisos solo aparecen y funcionan para Administradores. Un Coordinador no puede modificar sus propios permisos ni los de otra persona.",
)

add_heading(doc, "5. Comportamiento al cambiar un permiso", 1)
add_body(doc, "Para que una modificación sea segura y se refleje de inmediato, el flujo esperado es:")
permission_flow_num = create_decimal_numbering(doc)
for item in (
    "El Administrador activa o desactiva un permiso.",
    "El servidor valida que el actor sea Administrador y guarda el cambio antes de actualizar la interfaz.",
    "La aplicación emite un evento en tiempo real para las sesiones afectadas.",
    "La sesión del Coordinador vuelve a cargar sus permisos sin cerrar sesión.",
    "La navegación, botones y datos visibles se actualizan inmediatamente.",
    "Si el permiso fue retirado mientras la persona estaba dentro de una pantalla restringida, la aplicación cierra esa vista o redirige a una zona permitida.",
    "Aunque la interfaz todavía no se haya actualizado, el servidor rechaza cualquier nueva operación no autorizada.",
):
    add_number(doc, item, permission_flow_num)

add_heading(doc, "6. Reglas técnicas y de seguridad", 1)
for item in (
    "La base de datos y el servidor son la fuente de verdad. localStorage o cookies pueden ayudar con la experiencia, pero no deben autorizar acciones.",
    "Toda acción sensible usa una autorización central basada en: actor, acción, alcance y persona objetivo.",
    "Las páginas y consultas del Coordinador de comité siempre aplican el filtro del comité asignado.",
    "Los permisos configurables se guardan por tipo de Coordinador y, si se requiere, pueden admitir excepciones individuales futuras.",
    "Al revocar acceso, se invalida cualquier dato en caché que ya no corresponda al alcance del usuario.",
    "Las acciones de asistencia se validan en el servidor, incluso cuando se iniciaron desde un escáner QR.",
):
    add_bullet(doc, item)

add_heading(doc, "7. Auditoría requerida", 1)
add_body(doc, "El registro de auditoría debe identificar claramente quién realizó cada operación y sobre qué voluntario se aplicó.")
for item in (
    "Cambios de rol, tipo de Coordinador o permisos: valor anterior, valor nuevo y Administrador responsable.",
    "Escaneo QR, check-in y check-out: persona que realizó la operación, voluntario y hora efectiva.",
    "Registro de asistencia faltante y correcciones manuales: actor, motivo, valores anteriores y nuevos.",
    "Creación, importación, edición, archivo y reagendamiento de voluntarios.",
    "Resultado de la acción, incluidos intentos rechazados por falta de permiso cuando sean relevantes para seguridad.",
):
    add_bullet(doc, item)

add_heading(doc, "8. Migración desde el rol Editor", 1)
migration_num = create_decimal_numbering(doc)
for item in (
    "Cambiar el nombre visible “Editor” por “Coordinador”.",
    "Agregar un tipo de Coordinador: Comité o Tecnología.",
    "Migrar inicialmente a los Editores existentes como Coordinadores de comité para evitar otorgar acceso global por accidente.",
    "Un Administrador identifica manualmente a quienes pertenecen a Tecnología y cambia su tipo.",
    "Conservar compatibilidad temporal con el valor anterior mientras se actualizan sesiones y datos existentes.",
):
    add_number(doc, item, migration_num)

add_heading(doc, "9. Escenarios mínimos de aceptación", 1)
tests = [
    ("QR por Tecnología", "Un Coordinador de tecnología encuentra a una persona de cualquier comité y completa check-in/check-out."),
    ("QR bloqueado para Comité", "Un Coordinador de comité no ve el escáner y el servidor rechaza cualquier intento directo."),
    ("Alcance de Comité", "El Coordinador solo consulta y reagenda voluntarios de su comité."),
    ("Cambio inmediato", "Un Administrador apaga un permiso y la otra sesión pierde la acción sin volver a iniciar sesión."),
    ("Edición configurable", "Al apagar la edición personal para Comité, el perfil continúa visible pero no editable."),
    ("Asistencia configurable", "Tecnología solo registra entradas faltantes o corrige horarios cuando los permisos correspondientes están activos."),
    ("Administración protegida", "Ningún Coordinador puede cambiar roles, permisos o ajustes del sistema."),
    ("Auditoría completa", "Cada acción sensible muestra actor, voluntario, fecha y cambio realizado."),
]

table2 = doc.add_table(rows=1, cols=2)
headers2 = ("Escenario", "Resultado esperado")
for idx, text in enumerate(headers2):
    cell = table2.rows[0].cells[idx]
    set_cell_shading(cell, LIGHT)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0)
    add_text(p, text, bold=True, size=9, color=BLACK)
set_repeat_table_header(table2.rows[0])
for row_idx, row in enumerate(tests):
    cells = table2.add_row().cells
    for idx, text in enumerate(row):
        if row_idx % 2:
            set_cell_shading(cells[idx], "FAFAFA")
        set_cell_border(cells[idx])
        p = cells[idx].paragraphs[0]
        style_paragraph(p, after=0, line=1.05)
        add_text(p, text, size=9, bold=(idx == 0))
set_table_geometry(table2, [2880, 6480])

add_heading(doc, "10. Decisiones propuestas para confirmar", 1)
add_body(doc, "Estos puntos no cambian la separación principal entre roles, pero conviene confirmarlos antes de implementar:")
for item in (
    "¿Tecnología debe poder crear, importar o archivar voluntarios cuando un Administrador habilite esos permisos, o deben permanecer siempre exclusivos de Administradores?",
    "¿Tecnología debe tener activos por defecto avisos, solicitudes y reportes globales?",
    "¿El Coordinador de comité debe ver reportes limitados a su comité por defecto, o solo cuando un Administrador lo habilite?",
    "¿Se permitirán excepciones de permisos por persona además de la configuración general por tipo de Coordinador?",
):
    add_bullet(doc, item)

add_callout(
    doc,
    "Estado del documento",
    "Borrador para revisión. No autoriza cambios en producción hasta que se confirmen los puntos pendientes y la matriz final.",
)

# Document metadata
doc.core_properties.title = "Estructura de roles y permisos"
doc.core_properties.subject = "Definición funcional de roles, permisos y alcance para Volunteer Manager"
doc.core_properties.author = "Volunteer Manager"
doc.core_properties.keywords = "roles, permisos, administrador, coordinador de tecnología, coordinador de comité"

doc.save(OUTPUT)
print(OUTPUT)
